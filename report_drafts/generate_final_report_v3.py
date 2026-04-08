from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "report_drafts" / "final_report_draft_v3.docx"


def set_run_font(run, *, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_paragraph_style(paragraph, *, size=12, bold=False, italic=False, align=None, before=0, after=8, line_spacing=1.5):
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold or bool(run.bold), italic=italic or bool(run.italic))


def add_paragraph(doc, text, *, style=None, size=12, bold=False, italic=False, align=None, before=0, after=8, line_spacing=1.5):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    set_paragraph_style(
        p,
        size=size,
        bold=bold,
        italic=italic,
        align=align,
        before=before,
        after=after,
        line_spacing=line_spacing,
    )
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=15, bold=True)
        set_paragraph_style(p, size=15, bold=True, before=12, after=6)
        p.style = doc.styles["Heading 1"]
    elif level == 2:
        set_run_font(run, size=13, bold=True)
        set_paragraph_style(p, size=13, bold=True, before=8, after=4)
        p.style = doc.styles["Heading 2"]
    else:
        set_run_font(run, size=12, bold=True)
        set_paragraph_style(p, size=12, bold=True, before=6, after=3)
        p.style = doc.styles["Heading 3"]
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, rows, *, widths=None, header=True, font_size=11):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = width
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if header and r_idx == 0:
                shade_cell(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.15
                if header and r_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=header and r_idx == 0)
    return table


def add_figure_placeholder(doc, title, description, *, caption=None):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(6.3)

    header = table.cell(0, 0)
    header.text = f"Figure Placeholder: {title}"
    shade_cell(header, "EAF3E6")
    header.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in header.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, size=10.5, bold=True)

    body = table.cell(1, 0)
    body.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    body.text = description
    for paragraph in body.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=10.0, italic=True)

    if caption:
        add_paragraph(
            doc,
            f"Suggested caption: {caption}",
            size=10.0,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            before=2,
            after=8,
            line_spacing=1.0,
        )
    return table


def add_gallery_placeholder(doc, title, items, *, columns=2, caption=None):
    add_paragraph(
        doc,
        f"Figure Placeholder: {title}",
        size=10.5,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=6,
        after=4,
        line_spacing=1.0,
    )
    row_count = (len(items) + columns - 1) // columns
    table = doc.add_table(rows=row_count, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell_width = 6.3 / columns
    for col_idx in range(columns):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(cell_width)

    item_idx = 0
    for row_idx in range(row_count):
        for col_idx in range(columns):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if item_idx < len(items):
                cell.text = f"Insert image: {items[item_idx]}"
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(10)
                    paragraph.paragraph_format.space_after = Pt(10)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_run_font(run, size=10.0, italic=True)
            else:
                cell.text = ""
            item_idx += 1

    if caption:
        add_paragraph(
            doc,
            f"Suggested caption: {caption}",
            size=10.0,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            before=2,
            after=8,
            line_spacing=1.0,
        )
    return table


def add_field(paragraph, instruction, placeholder=""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    if placeholder:
        run = paragraph.add_run(placeholder)
        set_run_font(run, size=12, italic=True)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)

    for name, size in [("Title", 20), ("Subtitle", 14), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = name.startswith("Heading") or name == "Title"


def configure_footer(section):
    footer = section.footer
    p1 = footer.paragraphs[0]
    note = p1.add_run("AI-assisted draft; replace this with the final ENEL 500 AI disclosure before submission.")
    set_run_font(note, size=9, italic=True)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.line_spacing = 1.0

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p2, "PAGE", "1")
    for run in p2.runs:
        set_run_font(run, size=9)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0


def add_cover_page(doc):
    add_paragraph(
        doc,
        "EMG-Driven Driving Interface with Fixed Sensor Placement and CARLA-Based Evaluation",
        style="Title",
        size=20,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=36,
        after=12,
        line_spacing=1.2,
    )
    add_paragraph(
        doc,
        "Final Design Report Draft",
        style="Subtitle",
        size=14,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "ENEL 500 Capstone",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
        line_spacing=1.0,
    )

    cover_rows = [
        ["Item", "Draft entry"],
        ["Project title", "EMG-Driven Driving Interface with Fixed Sensor Placement and CARLA-Based Evaluation"],
        ["Team number", "[Insert team number]"],
        ["Project manager", "[Insert name, email, and UCID]"],
        ["Other team members", "[Insert names and UCIDs]"],
        ["Support technician", "[Insert if applicable, otherwise mark N/A]"],
        ["Academic advisor", "[Insert advisor name]"],
        ["Sponsoring company", "[Insert sponsor / lab / organization]"],
        ["Sponsor representative", "[Insert sponsor representative name and email]"],
        ["Final submission date", "[Replace with final submission date]"],
    ]
    add_table(doc, cover_rows, widths=[Inches(2.1), Inches(4.8)], font_size=10.5)

    add_paragraph(
        doc,
        f"This draft was updated against the maintained repository state on {date.today().isoformat()} and is intended to be finalized by inserting the final metric tables, figures, cost numbers, and team-specific reflections.",
        size=10.5,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=10,
        after=0,
        line_spacing=1.15,
    )
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u', "Right-click this field in Word and choose 'Update Field' before exporting the final PDF.")
    set_paragraph_style(p, italic=True, after=10)
    doc.add_page_break()


def add_glossary(doc):
    add_heading(doc, "Glossary", level=1)
    rows = [
        ["Term", "Definition"],
        ["EMG", "Electromyography; electrical activity measured from muscle contractions at the skin surface."],
        ["MVC", "Maximum voluntary contraction, used here as a high-effort calibration reference."],
        ["Strict layout", "The fixed pair-number sensor-placement contract enforced during training and live inference."],
        ["Per-subject model", "A model trained and evaluated for one participant to estimate personalized best-case performance."],
        ["Cross-subject model", "A pooled model evaluated with leave-one-subject-out testing to estimate generalization."],
        ["GestureCNNv2", "The maintained residual 1D CNN architecture used for EMG gesture classification."],
        ["Published gesture output", "The fused dual-arm realtime output exposed through get_latest_published_gestures()."],
        ["End-to-end latency", "The delay from the end of an EMG input window to the corresponding control application in CARLA."],
        ["Scenario success", "The scenario-level success/fail outcome logged by the CARLA evaluation runtime."],
    ]
    add_table(doc, rows, widths=[Inches(1.8), Inches(5.1)], font_size=10.5)


def add_ai_use_section(doc):
    add_heading(doc, "AI Use Disclosure and Verification", level=1)
    add_paragraph(
        doc,
        "This draft was prepared with AI-assisted writing and codebase cross-checking support. Before submission, the team should replace this draft disclosure with the exact ENEL 500-compliant statement listing every tool used, the tool or model version, dates of use, which report sections were assisted, whether any project data were exposed, the verification steps performed by the team, and the location of the maintained AI log.",
    )
    rows = [
        ["Required disclosure item", "Draft status / action needed before submission"],
        ["Tool name and version", "[Insert the exact assistant(s), model names, and version identifiers used during drafting and editing.]"],
        ["What the tool was used for", "Draft structuring, wording improvement, repo-aligned technical cross-checks, and document cleanup."],
        ["Project data exposure", "[State exactly what was shared with the tool, or confirm that no sensitive participant data were exposed.]"],
        ["Human verification steps", "All commands, metrics, file paths, and technical descriptions should be checked against the maintained repo state and final collected outputs."],
        ["AI log location", "[Insert the path or appendix location where the full AI interaction log is maintained.]"],
        ["Final compliance check", "[Confirm that the final disclosure matches the ENEL 500 policy sections required for the course.]"],
    ]
    add_table(doc, rows, widths=[Inches(2.3), Inches(4.6)], font_size=10.5)


def add_executive_summary(doc):
    add_heading(doc, "Executive Summary", level=1)
    add_paragraph(
        doc,
        "This capstone project developed an EMG-driven driving interface built around Delsys Trigno sensors, a fixed sensor-placement workflow, convolutional neural network classification, dual-arm realtime inference, and CARLA-based downstream control evaluation. The main engineering goal was to move from a broader but less reproducible prototype toward a technically defensible system whose data collection, preprocessing, runtime behavior, and evaluation outputs are all governed by an explicit contract.",
    )
    add_paragraph(
        doc,
        "The maintained pipeline begins with scripted four-gesture EMG collection through a custom Python GUI, followed by resampling to a shared 2000 Hz time grid and a matched offline filtering stack. GestureCNNv2 models are trained on overlapping EMG windows for both per-subject and cross-subject evaluation. Live deployment runs in dual-arm mode, uses the fixed gesture set neutral, left_turn, right_turn, and horn, and publishes fused split-or-single outputs that are consumed directly by the CARLA client.",
    )
    add_paragraph(
        doc,
        "The CARLA integration focuses the evaluation problem on a narrow but meaningful control contract. Steering is driven by the EMG inference output, reverse is triggered by a dual-horn event, and two named scenarios are used for controlled testing: lane_keep_5min and highway_overtake. Each evaluation run can automatically save both per-prediction realtime logs and per-tick drive logs, which are then converted into the final report metrics through the maintained evaluation scripts in eval_metrics/.",
    )
    add_paragraph(
        doc,
        "The final reported result groups for this design report are offline model metrics, end-to-end latency metrics, and CARLA scenario metrics. At the time of this draft revision, the system narrative, methods, and validation structure are aligned to the current repository scope. Final numeric values, figures, cost totals, and project-management reflections should be inserted after the remaining evaluation runs are completed and the metric tables are frozen.",
    )
    add_heading(doc, "Project Snapshot", level=2)
    rows = [
        ["Area", "Current maintained status", "What still needs to be finalized"],
        ["Data collection", "Scripted four-gesture strict-layout collection is implemented and documented.", "Insert final subject/session counts used in the reported experiments."],
        ["Preprocessing", "Resampling and filtering are implemented on the strict dataset roots.", "Confirm the exact datasets included in the final report tables."],
        ["Modeling", "GestureCNNv2 is maintained for per-subject and cross-subject training.", "Insert final offline metrics harvested from the selected bundles."],
        ["Realtime", "Dual-arm live inference is maintained with the current manual runtime tuning preset.", "Insert final latency results from logged evaluation runs."],
        ["Simulator evaluation", "Named CARLA scenarios and logging wrappers are implemented and documented.", "Insert final lane-keeping and overtake scenario results."],
        ["Report completion", "Narrative sections are aligned with the current repo scope.", "Add final cost data, finalized references, and team-specific reflection details."],
    ]
    add_table(doc, rows, widths=[Inches(1.5), Inches(2.8), Inches(2.6)], font_size=10.2)


def add_motivation_and_objectives(doc):
    add_heading(doc, "Project Motivation and Objectives", level=1)
    add_paragraph(
        doc,
        "The underlying motivation for the project is to evaluate whether surface EMG can support a practical driving-related control interface under controlled conditions. Traditional driving controls assume reliable mechanical interaction with a steering wheel, pedals, and other direct inputs. An EMG-based interface could provide an alternative control pathway for accessibility-oriented or non-traditional interaction settings, but only if the sensing and control stack is stable enough to operate beyond isolated classifier demos.",
    )
    add_paragraph(
        doc,
        "Earlier project iterations exposed three engineering risks that had to be resolved before the system could be evaluated credibly. First, inconsistent sensor placement weakened cross-session reproducibility because the same channel index did not always refer to the same physical muscle site. Second, mixed Delsys sensor families required explicit resampling and matched filtering so that training and realtime inference operated on comparable signals. Third, downstream simulator behavior needed to be measured with structured tasks and logged evidence rather than only anecdotal impressions or single-run screenshots.",
    )
    add_paragraph(
        doc,
        "The final project direction therefore prioritizes repeatability, auditability, and controlled validation over a broader gesture vocabulary or looser experimental scope. That design decision narrows the product story, but it produces a system that can be defended technically and reported transparently.",
    )
    rows = [
        ["Objective", "Engineering purpose", "How it will be demonstrated in the final report"],
        ["Enforce fixed sensor placement", "Keep channel meaning stable across sessions so model behavior remains interpretable.", "Describe the strict pair mapping and fail-closed layout checks."],
        ["Build a reproducible preprocessing path", "Align mixed-rate EMG channels before training and deployment.", "Report the 2000 Hz resampling and matched filtering workflow."],
        ["Train deployable CNN bundles", "Support both personalized and generalized model evaluation.", "Insert per-subject and cross-subject offline metrics from the maintained bundles."],
        ["Stabilize live dual-arm inference", "Publish usable gesture outputs for downstream control without relying on ad hoc manual interpretation.", "Report the runtime configuration and latency analysis results."],
        ["Demonstrate simulator-side control", "Show that the interface can support structured CARLA tasks with measurable control quality.", "Insert lane_keep_5min and highway_overtake results with scenario success and control-quality metrics."],
        ["Produce a reusable evaluation workflow", "Make future testing repeatable instead of spreadsheet-driven.", "Reference the maintained eval_metrics scripts, tables, and plots."],
    ]
    add_table(doc, rows, widths=[Inches(1.8), Inches(2.3), Inches(2.8)], font_size=10.2)


def add_methodology_and_design(doc):
    add_heading(doc, "Methodology and Design", level=1)
    add_heading(doc, "System Overview", level=2)
    add_paragraph(
        doc,
        "The maintained end-to-end pipeline is collection -> resampling -> filtering -> windowing and label generation -> CNN training -> dual-arm realtime inference -> CARLA control evaluation. The active collection entrypoint is DelsysPythonGUI.py. The main preprocessing entrypoints are emg/resample_raw_dataset.py and emg/filtering.py. Training is handled by train_per_subject.py and train_cross_subject.py. Standalone live inference is provided by realtime_gesture_cnn.py, and the canonical simulator client is carla_integration/manual_control_emg.py.",
    )
    rows = [
        ["Stage", "Maintained entrypoint(s)", "Main output"],
        ["Collection", "DelsysPythonGUI.py", "Raw strict-layout EMG sessions under data_strict/"],
        ["Resampling", "emg/resample_raw_dataset.py", "Common-rate EMG files on a 2000 Hz grid"],
        ["Filtering", "emg/filtering.py", "Filtered sessions for training and calibration-aware normalization"],
        ["Training", "train_per_subject.py / train_cross_subject.py", "GestureCNNv2 bundles under models/strict/"],
        ["Realtime inference", "realtime_gesture_cnn.py", "Published dual-arm gesture outputs and optional prediction logs"],
        ["CARLA evaluation", "carla_integration/manual_control_emg.py", "Per-tick drive logs and paired realtime logs for metric analysis"],
    ]
    add_table(doc, rows, widths=[Inches(1.3), Inches(2.4), Inches(3.2)], font_size=10.2)
    add_figure_placeholder(
        doc,
        "System pipeline and maintained workflow",
        "Insert a report-ready architecture figure showing collection, resampling, filtering, CNN training, dual-arm realtime inference, and CARLA evaluation.",
        caption="Maintained end-to-end workflow from Delsys collection to CARLA-based evaluation.",
    )
    add_paragraph(
        doc,
        "The placeholder above should be replaced with one architecture figure or flowchart showing the maintained pipeline. The table above captures the file-level entrypoints that should remain consistent with the final figure.",
        italic=True,
        size=10.5,
        after=10,
        line_spacing=1.15,
    )

    add_heading(doc, "Data Collection Pipeline", level=2)
    add_paragraph(
        doc,
        "Data collection is performed through DelsysPythonGUI.py and the underlying DataCollector/CollectDataWindow.py flow. The maintained collection workflow is fixed to a four-gesture protocol named standard_4g. The trainable labels are left_turn, right_turn, neutral, and horn. The collection GUI preserves per-channel metadata, including the EMG channel labels that later allow the system to resolve fixed pair identity rather than trusting scan order.",
    )
    add_paragraph(
        doc,
        "The active collection protocol uses 5.0 s prompted gesture intervals, 5.0 s neutral intervals, five repetitions, and optional neutral and MVC calibration captures. Inter-gesture rest is labeled neutral_buffer and intentionally trimmed out before training. This is important because the neutral_buffer segments capture transition or recovery behavior that should not be treated as clean steady-state gesture data.",
    )
    rows = [
        ["Protocol item", "Maintained setting"],
        ["Protocol name", "standard_4g"],
        ["Trainable labels", "left_turn, right_turn, neutral, horn"],
        ["Prompt duration", "5.0 s per gesture"],
        ["Neutral duration", "5.0 s"],
        ["Repetitions", "5"],
        ["Calibration", "Neutral-rest capture and MVC capture enabled by default"],
        ["Rest label", "neutral_buffer, removed before training"],
        ["Saved file type", "Compressed NPZ files with EMG arrays, timestamps, labels, events, metadata, and calibration arrays"],
    ]
    add_table(doc, rows, widths=[Inches(2.1), Inches(4.8)], font_size=10.5)
    add_gallery_placeholder(
        doc,
        "Four maintained gesture prompts",
        ["neutral", "left_turn", "right_turn", "horn"],
        columns=2,
        caption="Reference images for the four maintained gestures used in collection and deployment.",
    )

    add_heading(doc, "Strict Sensor Placement Workflow", level=2)
    add_paragraph(
        doc,
        "The strict sensor-placement policy is the most important systems-level design decision in the final maintained branch. Instead of accepting mixed historical layouts or attempting to salvage partially inconsistent sessions, the current workflow requires the same sensor pair identities to remain in the same physical positions across sessions. The helper logic in emg/strict_layout.py resolves channels from pair identity and fails closed when required pair numbers or expected arm-specific channel counts are missing.",
    )
    rows = [
        ["Arm", "Required pair order", "Expected channel count"],
        ["Right", "1, 2, 3, 7, 9, 11", "17"],
        ["Left", "4, 5, 6, 8, 10", "16"],
    ]
    add_table(doc, rows, widths=[Inches(1.2), Inches(3.0), Inches(1.8)], font_size=10.5)
    add_gallery_placeholder(
        doc,
        "Sensor placement reference photos",
        ["Right-arm sensor placement", "Left-arm sensor placement"],
        columns=2,
        caption="Reference sensor placement on the right and left arms for the strict-layout workflow.",
    )
    add_paragraph(
        doc,
        "This conservative policy is justified by the project objective: a reproducible deployment pipeline matters more than maximizing salvage of imperfect historical data. If the channel contract is unstable, any later offline score or CARLA metric becomes hard to interpret. Failing closed is therefore preferable to silently remapping channels and overstating reproducibility.",
    )

    add_heading(doc, "Preprocessing Pipeline", level=2)
    add_paragraph(
        doc,
        "After collection, strict sessions are resampled and filtered before they are used for training or final reporting. Resampling is required because the Delsys setup can include different sensor families with different effective sample rates. The maintained resampler estimates per-channel timing from timestamps, finds the overlapping valid time interval, constructs a common 2000.0 Hz grid, linearly interpolates each EMG channel onto that grid, and remaps labels by nearest-neighbor timing. Calibration arrays are also carried forward so live normalization can operate on signal data that have been processed with the same signal chain as the training windows.",
    )
    add_paragraph(
        doc,
        "The active filtering stack applies 60 Hz and 120 Hz notch filters followed by a 20-450 Hz sixth-order bandpass filter. The same broad filter contract is mirrored in the live path so the model sees comparable signal characteristics during training and deployment.",
    )
    rows = [
        ["Stage", "Current maintained behavior", "Why it matters"],
        ["Resampling", "Interpolate all channels onto a shared 2000 Hz grid", "Guarantees temporal alignment across mixed-rate sensors"],
        ["Label remapping", "Nearest-neighbor transfer onto the resampled grid", "Preserves the prompted label timing after interpolation"],
        ["Filtering", "60 Hz notch, 120 Hz notch, 20-450 Hz bandpass", "Suppresses power-line contamination and off-band noise"],
        ["Calibration carry-forward", "Neutral and MVC arrays are preserved through preprocessing", "Keeps training and runtime normalization auditable"],
    ]
    add_table(doc, rows, widths=[Inches(1.4), Inches(2.7), Inches(2.0)], font_size=10.0)

    add_heading(doc, "Windowing, Label Generation, and Model Architecture", level=2)
    add_paragraph(
        doc,
        "The maintained training path operates on filtered EMG windows rather than hand-engineered feature vectors. Windows are 200 samples long with a 100-sample step, which corresponds to a 100 ms window and 50 ms stride at the 2000 Hz target rate. Each window is labeled by majority vote over the contained sample labels, while neutral_buffer windows are dropped and low-purity windows can be excluded through the minimum label-confidence threshold.",
    )
    add_paragraph(
        doc,
        "The active classifier is GestureCNNv2 in emg/gesture_model_cnn.py. It is a residual 1D CNN that consumes input shaped as channels x time, applies input InstanceNorm1d, extracts temporal features across three backbone stages at 32, 64, and 128 channels, and uses squeeze-and-excitation channel attention inside each stage. A key design element is the explicit pre-normalization energy bypass, which preserves amplitude information that helps distinguish near-rest windows from active gesture windows even when per-window normalization would otherwise compress those cues.",
    )
    rows = [
        ["Component", "Maintained configuration"],
        ["Window size", "200 samples"],
        ["Window step", "100 samples"],
        ["Window labeling", "Majority vote over sample labels"],
        ["Removed label", "neutral_buffer"],
        ["Architecture", "GestureCNNv2 residual 1D CNN with channel attention"],
        ["Feature stages", "32, 64, 128"],
        ["Special feature", "Pre-normalization energy scalar concatenated into the classifier head"],
    ]
    add_table(doc, rows, widths=[Inches(2.0), Inches(4.9)], font_size=10.5)

    add_heading(doc, "Training Configuration", level=2)
    add_paragraph(
        doc,
        "The maintained branch keeps both a per-subject and a cross-subject training path because they answer different validation questions. Per-subject models estimate personalized best-case performance under repeated collection from the same user. Cross-subject models estimate whether the system can generalize to a held-out participant. The training bundle metadata preserve label maps, architecture settings, training settings, stored metrics, and sensor-layout information so the selected model artifacts remain auditable later in the report.",
    )
    rows = [
        ["Setting", "Per-subject training", "Cross-subject training"],
        ["Window size / step", "200 / 100 samples", "200 / 100 samples"],
        ["Epochs", "40", "50"],
        ["Batch size", "512", "512"],
        ["Optimizer", "Adam", "Adam"],
        ["Learning rate", "1e-4", "1e-4"],
        ["Dropout", "0.25", "0.25"],
        ["Label smoothing", "0.05", "0.05"],
        ["Minimum label confidence", "0.85", "0.85 on the maintained four-gesture path"],
        ["Main validation split", "Session-grouped StratifiedGroupKFold", "Leave-one-subject-out evaluation before the pooled final fit"],
        ["Sampling emphasis", "Session separation to prevent leakage", "WeightedRandomSampler for subject-balanced sampling"],
    ]
    add_table(doc, rows, widths=[Inches(1.8), Inches(2.5), Inches(2.5)], font_size=9.8)

    add_heading(doc, "Realtime Inference and Tuning", level=2)
    add_paragraph(
        doc,
        "The active realtime entrypoint is realtime_gesture_cnn.py. The maintained deployment path is dual-arm only and is fixed to the four-gesture set neutral, left_turn, right_turn, and horn. At runtime, the script resolves fixed channel order from live Delsys metadata, resamples live input streams to 2000 Hz, mirrors the maintained filtering stack, performs live neutral and MVC calibration when calibration quality is sufficient, and then publishes a fused dual-arm output through get_latest_published_gestures(). That published interface is the source of truth for downstream control behavior.",
    )
    add_paragraph(
        doc,
        "The current checked-in runtime preset is the manual preset in emg/runtime_tuning.py. It uses smoothing of 1 frame, a minimum confidence threshold of 0.80, a dual-arm agreement threshold of 0.55, no output hysteresis, and no softmax-margin rejection. These settings should be reported as the baseline runtime configuration unless the final experiments intentionally freeze a different preset and that change is documented explicitly.",
    )
    rows = [
        ["Runtime item", "Current maintained value / behavior"],
        ["Mode", "Dual-arm live inference"],
        ["Active gestures", "neutral, left_turn, right_turn, horn"],
        ["Resampled rate", "2000 Hz"],
        ["Windowing", "200-sample windows with 100-sample step"],
        ["Preset name", "manual"],
        ["Smoothing", "1 frame"],
        ["Minimum confidence", "0.80"],
        ["Dual-arm agreement threshold", "0.55"],
        ["Output hysteresis", "Disabled"],
        ["Softmax-margin rejection", "Disabled"],
        ["Published interface", "get_latest_published_gestures()"],
    ]
    add_table(doc, rows, widths=[Inches(2.2), Inches(4.7)], font_size=10.0)

    add_heading(doc, "CARLA Integration", level=2)
    add_paragraph(
        doc,
        "The canonical simulator client is carla_integration/manual_control_emg.py. The preferred evaluation path is keyboard-first plus EMG steering control, although the current client still retains optional joystick-based throttle and brake support when one joystick is connected. The EMG contract is intentionally narrow: the left arm issues stronger steering commands, the right arm issues lighter steering commands, and a dual-horn event requests reverse toggling. That narrow contract keeps the evaluation problem focused on controllable steering behavior instead of attempting a full EMG-only longitudinal controller.",
    )
    add_paragraph(
        doc,
        "Two named scenario wrappers are maintained for evaluation: lane_keep_5min and highway_overtake. Both run on Town04_Opt. Free-roam practice still defaults to Town03_Opt with 90 ambient vehicles and 0 pedestrians. The client retains the inherited Tab camera toggle, but the preferred evaluation view is the standard RGB camera. Each named evaluation wrapper launches the CARLA client with --eval-log-dir so the run automatically saves timestamped carla_drive_<timestamp>.csv and realtime_predictions_<timestamp>.csv files for later analysis.",
    )
    rows = [
        ["Scenario / mode", "Map", "Purpose", "Primary success condition"],
        ["Free roam practice", "Town03_Opt", "Driver familiarization and informal tuning", "No formal success metric; used only for practice"],
        ["lane_keep_5min", "Town04_Opt", "Structured lane-keeping route", "Reach the scenario finish condition without relying on informal interpretation"],
        ["highway_overtake", "Town04_Opt", "Structured pass-and-return task", "Complete the overtake objective and cross the finish condition"],
    ]
    add_table(doc, rows, widths=[Inches(1.9), Inches(1.1), Inches(1.8), Inches(2.1)], font_size=9.8)
    add_figure_placeholder(
        doc,
        "Driving rig and operator setup",
        "Insert one photo showing the workstation, input posture, and simulator display used during EMG driving tests.",
        caption="Physical testing rig used for realtime EMG-controlled CARLA evaluation.",
    )


def add_scope_and_functionality(doc):
    add_heading(doc, "Product Scope and Final Product Functionality", level=1)
    add_paragraph(
        doc,
        "The delivered product is a working end-to-end EMG control research platform rather than a consumer-ready assistive driving system. In its final maintained form, it supports strict-layout EMG collection, fixed-root preprocessing, per-subject and cross-subject CNN training, dual-arm realtime inference on a fixed four-gesture set, and CARLA-based downstream evaluation with repeatable logging. That scope is sufficient for the capstone objective because it allows the team to move from raw acquisition to reportable control metrics within one coherent workflow.",
    )
    add_paragraph(
        doc,
        "The final scope is intentionally narrower than the broader exploratory direction that existed earlier in the project. The current report should frame that narrowing as an engineering choice rather than a limitation hidden after the fact. A smaller, stable, and auditable interface is more valuable at this stage than a broader interface whose data contract or evaluation logic cannot be defended consistently.",
    )
    rows = [
        ["Area", "Included in the final delivered scope", "Not claimed / intentionally out of scope"],
        ["Gesture vocabulary", "Fixed four-gesture workflow: neutral, left_turn, right_turn, horn", "Broader exploratory label sets are not part of the maintained final story"],
        ["Sensor handling", "Fixed pair-number layout with fail-closed validation", "Automatic recovery from arbitrary mixed layouts"],
        ["Realtime behavior", "Dual-arm publication for downstream control", "A claim of universally robust real-world deployment across all users"],
        ["Vehicle control", "Steering and reverse logic driven by EMG in CARLA", "Full EMG-only throttle and brake control as a maintained report claim"],
        ["Validation", "Offline, latency, and scenario-level CARLA metrics", "Unstructured anecdotal testing as primary evidence"],
        ["Deployment environment", "Simulator evaluation in CARLA", "On-road vehicle deployment or safety certification"],
    ]
    add_table(doc, rows, widths=[Inches(1.5), Inches(2.7), Inches(2.7)], font_size=9.8)


def add_technical_specs(doc):
    add_heading(doc, "Technical Specifications of the Final Product", level=1)
    add_paragraph(
        doc,
        "This section records the maintained technical specifications of the final system rather than the broader exploratory options that existed earlier in development. If the final submission differs from any Fall-term target, that change should be described as a deliberate scope correction in support of reproducibility, auditability, and metric-driven validation.",
    )
    add_heading(doc, "Acquisition and Data Specification", level=2)
    rows = [
        ["Specification", "Final maintained value"],
        ["Data roots", "data_strict/, data_resampled_strict/, models/strict/"],
        ["Right-arm channel count", "17 strict channels"],
        ["Left-arm channel count", "16 strict channels"],
        ["Target resampled rate", "2000 Hz"],
        ["Training labels", "left_turn, right_turn, neutral, horn"],
        ["Removed label", "neutral_buffer is a collection artifact, not a trainable class"],
        ["Windowing", "200-sample windows, 100-sample step"],
        ["Filtering", "60 Hz notch, 120 Hz notch, 20-450 Hz bandpass"],
    ]
    add_table(doc, rows, widths=[Inches(2.3), Inches(4.6)], font_size=10.5)

    add_heading(doc, "Model and Runtime Specification", level=2)
    rows = [
        ["Specification", "Final maintained value"],
        ["Architecture", "GestureCNNv2 residual 1D CNN with channel attention and energy bypass"],
        ["Per-subject output pattern", "models/strict/per_subject/<arm>/<subject>v6_4_gestures.pt"],
        ["Cross-subject output pattern", "models/strict/cross_subject/<arm>/v6_4_gestures.pt"],
        ["Checked-in realtime defaults", "Matthew right-arm and left-arm per-subject bundles"],
        ["Runtime preset", "manual"],
        ["Realtime minimum confidence", "0.80"],
        ["Dual-arm agreement threshold", "0.55"],
        ["Output hysteresis", "Disabled in the current maintained preset"],
        ["Softmax rejection", "Disabled in the current maintained preset"],
    ]
    add_table(doc, rows, widths=[Inches(2.3), Inches(4.6)], font_size=10.2)

    add_heading(doc, "Simulator and Evaluation Specification", level=2)
    rows = [
        ["Specification", "Final maintained value"],
        ["Canonical CARLA client", "carla_integration/manual_control_emg.py"],
        ["Named evaluation scenarios", "lane_keep_5min and highway_overtake"],
        ["Scenario maps", "Town04_Opt for both named scenarios"],
        ["Free-roam practice", "Town03_Opt with 90 ambient vehicles and 0 pedestrians"],
        ["Eval wrappers", "lane_keep_5min_eval.cmd and highway_overtake_eval.cmd"],
        ["Auto-saved raw logs", "carla_drive_<timestamp>.csv and realtime_predictions_<timestamp>.csv"],
        ["Maintained metric families", "Offline model metrics, end-to-end latency, and CARLA scenario metrics"],
    ]
    add_table(doc, rows, widths=[Inches(2.3), Inches(4.6)], font_size=10.2)


def add_validation_section(doc):
    add_heading(doc, "Measuring Success and Validation Test Results", level=1)
    add_heading(doc, "Success Criteria", level=2)
    add_paragraph(
        doc,
        "The project should be judged successful if it demonstrates a reproducible fixed-layout EMG pipeline, deployable four-gesture realtime inference, traceable evaluation logging, and simulator-side task performance that can be reported with quantitative evidence. The report should emphasize that system success is not defined by one offline accuracy number alone. It depends on the coherence of the full acquisition-to-control chain.",
    )
    rows = [
        ["Success criterion", "Why it matters", "Evidence source in the final report"],
        ["Stable sensor-layout contract", "Without stable channel meaning, no later metric is fully trustworthy.", "Strict layout description, channel counts, and fail-closed checks"],
        ["Adequate offline classification quality", "The base classifier must distinguish the four maintained gestures consistently.", "Offline metrics harvested from the selected model bundles"],
        ["Acceptable live timing", "The control loop must react quickly enough for practical simulator use.", "End-to-end latency summaries from paired realtime and CARLA logs"],
        ["Usable simulator control", "The user must be able to complete structured CARLA tasks with measurable control quality.", "Scenario success, completion time, and driving-metric summaries"],
        ["Repeatable evaluation workflow", "The report should be reproducible without spreadsheet-only manual processing.", "Documented eval_metrics pipeline and saved run artifacts"],
    ]
    add_table(doc, rows, widths=[Inches(1.9), Inches(2.2), Inches(2.8)], font_size=9.8)

    add_heading(doc, "Main Metric Set", level=2)
    add_paragraph(
        doc,
        "The final reported metric set in this design report should stay aligned with the maintained evaluation pipeline and with the research-paper draft. The primary groups are offline model metrics, end-to-end latency, and CARLA scenario metrics. Additional diagnostics such as mean_velocity_deviation_mps or command_success_rate can be mentioned when helpful, but they should not replace the headline metrics listed below.",
    )
    rows = [
        ["Metric group", "Required final reported metrics"],
        ["Offline model metrics", "balanced_accuracy, macro_precision, macro_recall, macro_f1, worst_class_recall"],
        ["Latency metrics", "mean end-to-end latency, p95 end-to-end latency"],
        ["CARLA metrics", "scenario success, completion time, mean velocity, lane offset mean, steering angle mean, steering entropy, lane error RMSE, lane invasions"],
    ]
    add_table(doc, rows, widths=[Inches(2.0), Inches(4.9)], font_size=10.2)

    add_heading(doc, "Validation Workflow", level=2)
    add_paragraph(
        doc,
        "Offline model metrics are harvested from saved bundles with eval_metrics/harvest_model_metrics.py. Live evaluation runs are collected by launching lane_keep_5min_eval.cmd or highway_overtake_eval.cmd, which automatically save timestamped CARLA and realtime logs. End-to-end latency is summarized by eval_metrics/analyze_latency.py, CARLA drive behavior is summarized by eval_metrics/analyze_drive_metrics.py, and participant-level plus aggregate tables are assembled by eval_metrics/build_eval_tables.py. The current plot layer supports plot_model_summary.py, plot_latency_summary.py, and plot_carla_summary_bars.py for the main report figures.",
    )
    add_paragraph(
        doc,
        "This report draft is written so that the final numbers can be inserted directly from those generated outputs. The narrative should not be finalized until the metric tables below are completed and cross-checked against the actual files placed under eval_metrics/out/.",
        italic=True,
        size=10.5,
        after=8,
        line_spacing=1.15,
    )

    add_heading(doc, "Results Placeholders", level=2)
    add_paragraph(
        doc,
        "Insert the final quantitative results into the following tables after the remaining evaluation runs are complete. If multiple conditions are reported, add or remove rows as needed, but keep the same metric columns so the final report remains aligned with the maintained evaluation scripts.",
    )

    add_heading(doc, "Offline Model Results", level=3)
    offline_rows = [
        ["Condition", "Arm / scope", "Balanced accuracy", "Macro precision", "Macro recall", "Macro F1", "Worst-class recall", "Notes"],
        ["[Insert condition]", "[Insert arm or pooled scope]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Bundle path or subject note]"],
        ["[Insert condition]", "[Insert arm or pooled scope]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Bundle path or subject note]"],
    ]
    add_table(doc, offline_rows, widths=[Inches(1.1), Inches(1.0), Inches(0.95), Inches(0.9), Inches(0.9), Inches(0.8), Inches(1.0), Inches(1.25)], font_size=9.2)
    add_figure_placeholder(
        doc,
        "Offline model summary plot",
        "Insert the figure exported by eval_metrics/plot_model_summary.py after the final model metrics are frozen.",
        caption="Offline model performance across the maintained metric set.",
    )

    add_heading(doc, "Latency Results", level=3)
    latency_rows = [
        ["Scenario / condition", "Participant or cohort", "Mean e2e latency (ms)", "P95 e2e latency (ms)", "Notes"],
        ["lane_keep_5min", "[Insert]", "[Insert]", "[Insert]", "[Timestamp range or run note]"],
        ["highway_overtake", "[Insert]", "[Insert]", "[Insert]", "[Timestamp range or run note]"],
    ]
    add_table(doc, latency_rows, widths=[Inches(1.6), Inches(1.4), Inches(1.3), Inches(1.3), Inches(1.3)], font_size=9.5)
    add_figure_placeholder(
        doc,
        "Latency summary plot",
        "Insert the figure exported by eval_metrics/plot_latency_summary.py after the final latency summaries are generated.",
        caption="Mean and p95 end-to-end latency for the maintained evaluation scenarios.",
    )

    add_heading(doc, "CARLA Scenario Results", level=3)
    carla_rows = [
        ["Scenario", "Scenario success", "Completion time (s)", "Mean velocity (m/s)", "Lane offset mean (m)", "Steering angle mean (rad)", "Steering entropy", "Lane error RMSE (m)", "Lane invasions"],
        ["lane_keep_5min", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]"],
        ["highway_overtake", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]", "[Insert]"],
    ]
    add_table(doc, carla_rows, widths=[Inches(1.15), Inches(0.9), Inches(0.85), Inches(0.85), Inches(0.85), Inches(0.95), Inches(0.8), Inches(0.85), Inches(0.8)], font_size=8.6)
    add_figure_placeholder(
        doc,
        "CARLA metric summary plot",
        "Insert the figure exported by eval_metrics/plot_carla_summary_bars.py after the aggregate CARLA tables are finalized.",
        caption="CARLA scenario performance across the maintained driving metric set.",
    )

    add_heading(doc, "Recommended Final Figures", level=3)
    figure_rows = [
        ["Figure", "Source", "Purpose in the final report"],
        ["System workflow figure", "Manual diagram", "Shows the maintained acquisition-to-evaluation pipeline at a glance"],
        ["Gesture reference collage", "Manual photo set", "Makes the four-label contract easy for the reader to interpret"],
        ["Sensor placement reference photos", "Manual photo set", "Documents the strict-layout contract visually"],
        ["Driving rig photo", "Manual photo", "Shows the actual operator and workstation setup used during testing"],
        ["Offline model summary figure", "eval_metrics/plot_model_summary.py", "Shows the maintained offline metric set in one report-ready comparison figure"],
        ["Latency summary figure", "eval_metrics/plot_latency_summary.py", "Summarizes mean and p95 end-to-end latency"],
        ["CARLA summary figure", "eval_metrics/plot_carla_summary_bars.py", "Summarizes the maintained CARLA metric group"],
        ["Representative CARLA screenshot", "Manual screenshot", "Optional visual context for the lane-keep or overtake scenario"],
    ]
    add_table(doc, figure_rows, widths=[Inches(1.8), Inches(2.0), Inches(3.1)], font_size=9.6)

    add_heading(doc, "Validation Discussion", level=2)
    add_paragraph(
        doc,
        "The final discussion in this section should interpret whether the fixed-layout redesign, the narrow four-gesture deployment contract, and the structured CARLA evaluation path materially improved system repeatability. In particular, the final write-up should connect the collected metric tables to the design choices described earlier: layout enforcement, mirrored preprocessing, dual-arm fusion, the maintained runtime thresholds, and explicit scenario logging. Claims about usability or feasibility should be tied directly to the measured results rather than to anecdotal driving impressions alone.",
    )


def add_cost_and_reflections(doc):
    add_heading(doc, "List of Tools, Materials, Supplies, Costs, etc.", level=1)
    add_paragraph(
        doc,
        "The final report should include a complete actual-versus-estimated cost view rather than only a general resource list. If certain resources were borrowed, shared, or already available through a lab or course environment, that should be stated clearly so the reader can distinguish direct project spending from infrastructure access.",
    )
    cost_rows = [
        ["Item", "Type", "Fall estimate", "Actual cost", "Variance / notes"],
        ["Delsys Trigno hardware access", "Equipment", "[Insert]", "[Insert]", "[Owned, borrowed, or purchased details]"],
        ["Windows workstation / GPU-capable host", "Compute", "[Insert]", "[Insert]", "[State whether this was existing infrastructure]"],
        ["CARLA simulator environment", "Software", "[Insert]", "[Insert]", "[If zero direct cost, state that clearly]"],
        ["Adapters, electrodes, or consumables", "Supplies", "[Insert]", "[Insert]", "[Insert variance reason]"],
        ["Any software licenses or subscriptions", "Services", "[Insert]", "[Insert]", "[Insert variance reason]"],
    ]
    add_table(doc, cost_rows, widths=[Inches(1.7), Inches(0.9), Inches(1.0), Inches(1.0), Inches(2.3)], font_size=9.4)

    add_heading(doc, "Reflection on Ethics and Equity in the Capstone Journey", level=1)
    add_paragraph(
        doc,
        "The strongest ethical issue in this project is the risk of overstating the readiness of a biosignal-based control interface for safety-critical use. Even though the system is evaluated in simulation, the language used in the report should remain careful and bounded. The responsible claim is that the project explores the feasibility of a constrained EMG control interface under simulator conditions, not that it establishes a deployable medical or automotive product. That distinction matters because end users could be misled if prototype performance is described without the surrounding limitations on subject count, sensor placement discipline, and controlled test conditions.",
    )
    add_paragraph(
        doc,
        "Equity considerations also arise in both the design and the project process. On the design side, an accessibility-motivated interface must still account for who is represented in the collected data and whose physiology or fatigue patterns may be underrepresented. On the team side, the project benefited from making expectations explicit: shared documentation, logged evaluation procedures, and transparent task boundaries reduce the chance that key knowledge stays with only one person. The final submission should add one or two concrete team-specific examples showing how ethical caution and equity awareness shaped real decisions during the capstone.",
    )

    add_heading(doc, "Reflection on Impact of Engineering on Society and Environment", level=1)
    add_paragraph(
        doc,
        "The positive societal case for this project is that it contributes to the broader study of alternative human-machine interfaces that may eventually support accessibility-oriented control schemes. The strongest value at this stage is not direct deployment, but the disciplined investigation of what makes such interfaces reliable or unreliable. That knowledge can help future work avoid overclaiming and focus on the engineering factors that genuinely matter for safety and usability.",
    )
    add_paragraph(
        doc,
        "The negative societal risk is that even promising prototype metrics could be interpreted too broadly outside their intended context. A simulator-based capstone cannot capture every failure mode that would matter in a real vehicle, and the report should state that clearly. Environmental impact is modest at the prototype stage and is mostly associated with hardware use, compute, and testing infrastructure. A positive offset is that simulator-heavy iteration can reduce unnecessary physical prototyping and allows more controlled validation before any riskier deployment setting is considered.",
    )

    add_heading(doc, "Reflection on the Team's Approach to Project Management", level=1)
    add_paragraph(
        doc,
        "One of the clearest project-management lessons in this capstone is that schedule pressure can encourage teams to keep widening scope even when the measurement pipeline is not ready. The more effective decision in this project was to narrow the maintained branch around a stable four-gesture workflow, strict dataset roots, and repeatable simulator scenarios. That shift turned the project from a loose set of promising scripts into a system that can actually be validated and reported coherently.",
    )
    add_paragraph(
        doc,
        "The final submission should personalize this section with the actual team schedule, milestone adjustments, communication practices, and risk-management steps used through the term. Recommended items to address are: how responsibilities were divided, how technical debt or stale experimental branches were handled, what the major schedule risks were, and which management decisions most improved momentum once the project scope was tightened.",
    )

    add_heading(doc, "Reflection on the Team's Approach to Economics and Cost Control", level=1)
    add_paragraph(
        doc,
        "The main economic lesson from this project is that engineering time and experimental ambiguity were more expensive than adding another exploratory feature. Preserving strict data roots, building reusable evaluation scripts, and narrowing the deployment claim all increase short-term effort, but they reduce the long-term cost of rerunning experiments whose assumptions were unclear. That is a valid form of cost control: spending effort early to avoid paying for low-confidence results later.",
    )
    add_paragraph(
        doc,
        "This section should be finalized with the actual budget comparison from the Fall term and with a short explanation of which design or scope decisions were driven by cost considerations. If existing lab equipment or borrowed resources were essential to feasibility, that should be stated explicitly because it affects the practical economic interpretation of the final design.",
    )

    add_heading(doc, "Reflection on Team Members' Approach to Life-Long Learning", level=1)
    add_paragraph(
        doc,
        "This project required learning across several domains at once: biosignal acquisition, signal preprocessing, neural network training, realtime systems, simulator integration, and evidence-based validation. A major professional takeaway is that good engineering often depends less on adding complexity and more on tightening the contract between system components. Another is that documentation and evaluation infrastructure are not secondary tasks; they are part of the actual technical solution when a project depends on repeatability.",
    )
    add_paragraph(
        doc,
        "The final report should personalize this section with team-specific growth examples, such as learning to interpret EMG quality, debug live timing issues, use structured validation splits, or redesign a workflow when the original assumptions were too loose. Those examples will make the reflection section stronger than a generic statement about learning new tools.",
    )


def add_conclusion_appendix_references(doc):
    add_heading(doc, "Conclusion and Recommendations", level=1)
    add_paragraph(
        doc,
        "The final delivered capstone system is a technically coherent EMG control pipeline that connects fixed-layout Delsys collection, reproducible preprocessing, residual CNN-based gesture classification, dual-arm realtime inference, and CARLA-based downstream evaluation. The most important engineering outcome is not merely that the system can produce predictions, but that the system now exposes enough structure to be measured and discussed responsibly.",
    )
    add_paragraph(
        doc,
        "The clearest next step is to freeze the final experiment set, insert the completed offline, latency, and CARLA results, and then revise the claims so each conclusion maps directly to measured evidence. If future work continues beyond the capstone, the recommended priority is to deepen data coverage and validation on the maintained four-gesture path before broadening the control vocabulary or making stronger deployment claims.",
    )

    add_heading(doc, "Appendix", level=1)
    add_heading(doc, "Appendix A. Metric Collection and Table-Building Commands", level=2)
    add_paragraph(
        doc,
        "These commands reflect the current maintained evaluation workflow and can be copied into the final report appendix once the exact file names and output locations are confirmed for the final experiment set.",
    )
    command_rows = [
        ["Step", "Command", "Purpose"],
        ["Collect lane-keep run", r"carla_integration\lane_keep_5min_eval.cmd", "Launches CARLA with automatic eval logging to eval_metrics/out/lane_keep_eval"],
        ["Collect overtake run", r"carla_integration\highway_overtake_eval.cmd", "Launches CARLA with automatic eval logging to eval_metrics/out/highway_overtake_eval"],
        ["Harvest offline metrics", r"python eval_metrics/harvest_model_metrics.py --models-root models/strict --output-csv eval_metrics/out/model_metrics.csv --output-json eval_metrics/out/model_metrics.json", "Exports bundle-level offline metrics"],
        ["Convenience post-run analysis", r"python eval_metrics/gather_current_metrics.py", "Stages discovered latency and drive summaries under eval_metrics/out/current_metrics/"],
        ["Manual latency summary", r"python eval_metrics/analyze_latency.py --realtime-log <realtime_csv> --carla-log <carla_csv> --output-json <latency_json> --output-csv <latency_joined_csv>", "Computes mean and p95 end-to-end latency"],
        ["Manual drive summary", r"python eval_metrics/analyze_drive_metrics.py --log <carla_csv> --output-json <drive_summary_json>", "Computes CARLA scenario metrics"],
        ["Build manifest", r"python eval_metrics/build_eval_tables.py --write-template-manifest eval_metrics/table_manifest.csv", "Creates a table-manifest starter file"],
        ["Build final tables", r"python eval_metrics/build_eval_tables.py --manifest eval_metrics/table_manifest.csv --model-metrics eval_metrics/out/model_metrics.csv --output-dir eval_metrics/out/tables", "Builds participant and aggregate report tables"],
        ["Offline model plot", r"python eval_metrics/plot_model_summary.py --input-csv eval_metrics/out/model_metrics.csv --gesture-bucket 4_gesture --latest-only", "Generates the model summary figure"],
        ["Latency plot", r"python eval_metrics/plot_latency_summary.py --input-csv eval_metrics/out/tables/evaluation_aggregate_table.csv --deliverable-bucket capstone_report", "Generates the latency summary figure"],
        ["CARLA plot", r"python eval_metrics/plot_carla_summary_bars.py --input-csv eval_metrics/out/tables/evaluation_aggregate_table.csv --deliverable-bucket capstone_report", "Generates the CARLA summary figure"],
    ]
    add_table(doc, command_rows, widths=[Inches(1.2), Inches(3.9), Inches(1.8)], font_size=8.8)

    add_heading(doc, "Appendix B. Suggested Final Figure and Table Inventory", level=2)
    appendix_rows = [
        ["Item", "Recommended status before submission"],
        ["Main offline results table", "Required"],
        ["Main latency results table", "Required"],
        ["Main CARLA scenario table", "Required"],
        ["System workflow figure", "Recommended"],
        ["Gesture reference collage", "Recommended"],
        ["Right-arm and left-arm sensor placement photos", "Recommended"],
        ["Driving rig photo", "Recommended"],
        ["Offline model summary figure", "Recommended"],
        ["Latency summary figure", "Recommended"],
        ["CARLA summary figure", "Recommended"],
        ["Representative lane-keep or overtake scenario screenshot", "Optional but useful if space allows"],
        ["Selected confusion matrix or supporting figure", "Optional appendix item"],
        ["Sample EMG trace or raw-vs-filtered comparison", "Optional appendix item"],
        ["Collection GUI screenshot", "Optional appendix item"],
    ]
    add_table(doc, appendix_rows, widths=[Inches(2.6), Inches(4.3)], font_size=10.0)

    add_heading(doc, "References", level=1)
    add_paragraph(
        doc,
        "[1] P. Basnet et al., \"Evaluating the Feasibility of EMG-Based Human-Machine Interfaces for Driving,\" Human Factors, vol. 68, no. 1, pp. 123-141, 2026, doi: 10.1177/00187208251367179.",
        after=2,
    )
    add_paragraph(
        doc,
        "[2] CARLA Simulator documentation for the version used in this project. Insert the complete citation or documentation URL and access date.",
        after=2,
    )
    add_paragraph(
        doc,
        "[3] Delsys Trigno hardware and SDK documentation used for acquisition and integration. Insert the complete citation details used by the team.",
        after=2,
    )
    add_paragraph(
        doc,
        "[4] Add any additional EMG, CNN, assistive-interface, or methodology references cited in the final text, using the citation style required by the course.",
        after=2,
    )


def main():
    doc = Document()
    configure_styles(doc)
    configure_footer(doc.sections[0])

    add_cover_page(doc)
    add_toc(doc)
    add_glossary(doc)
    add_ai_use_section(doc)
    add_executive_summary(doc)
    add_motivation_and_objectives(doc)
    add_methodology_and_design(doc)
    add_scope_and_functionality(doc)
    add_technical_specs(doc)
    add_validation_section(doc)
    add_cost_and_reflections(doc)
    add_conclusion_appendix_references(doc)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
